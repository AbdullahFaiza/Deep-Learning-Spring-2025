
import os
import re
import wikipedia
import gradio as gr
from urllib.parse import urlparse
from newspaper import Article
from serpapi import GoogleSearch
from transformers import pipeline
from docx import Document
from reportlab.pdfgen import canvas

# Initialize Hugging Face Summarization Pipeline
summarizer = pipeline("summarization")

# Function to perform a Google search and retrieve article links
def google_search(query):
    search = GoogleSearch({"q": query, "api_key": "your_serp_api_key"})
    results = search.get_dict()
    return results['organic_results']

# Function to fetch article content using Newspaper3k
def fetch_article(url):
    article = Article(url)
    article.download()
    article.parse()
    return article.text

# Function to summarize content using Hugging Face summarizer
def summarize_content(content):
    summary = summarizer(content, max_length=150, min_length=50, do_sample=False)
    return summary[0]['summary_text']

# Function to create a Word Document report
def create_word_report(content, summary, citations):
    doc = Document()
    doc.add_heading('Research Report', 0)
    doc.add_heading('Summary:', level=1)
    doc.add_paragraph(summary)
    
    doc.add_heading('Content:', level=1)
    doc.add_paragraph(content)
    
    doc.add_heading('Citations:', level=1)
    doc.add_paragraph(citations)
    
    doc.save("research_report.docx")
    return "research_report.docx"

# Function to create a PDF report
def create_pdf_report(content, summary, citations):
    pdf_file = "research_report.pdf"
    c = canvas.Canvas(pdf_file)
    c.drawString(100, 800, "Research Report")
    
    c.drawString(100, 780, "Summary:")
    c.drawString(100, 760, summary)
    
    c.drawString(100, 740, "Content:")
    c.drawString(100, 720, content)
    
    c.drawString(100, 700, "Citations:")
    c.drawString(100, 680, citations)
    
    c.save()
    return pdf_file

# Gradio interface function
def research_agent(query):
    # Perform Google search
    results = google_search(query)
    article_urls = [result['link'] for result in results]
    
    # Fetch the first article
    article_content = fetch_article(article_urls[0])
    
    # Summarize the article content
    summary = summarize_content(article_content)
    
    # Prepare citations (simple placeholder for now)
    citations = "\n".join([urlparse(url).netloc for url in article_urls])
    
    # Generate reports
    word_report = create_word_report(article_content, summary, citations)
    pdf_report = create_pdf_report(article_content, summary, citations)
    
    return word_report, pdf_report

# Gradio Interface Setup
interface = gr.Interface(fn=research_agent, inputs="text", outputs=["file", "file"])

if __name__ == "__main__":
    interface.launch()
